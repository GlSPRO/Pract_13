from __future__ import annotations

import csv
import sqlite3
from pathlib import Path

from docx import Document


ROOT = Path(__file__).resolve().parents[1]
DB_PATH = ROOT / "db.sqlite3"
DOCX_PATH = ROOT / "data_dictionary_one_table_ru.docx"
TSV_PATH = ROOT / "data_dictionary_one_table_ru.tsv"


FIELD_DESCRIPTIONS = {
    "id": "Идентификатор записи",
    "username": "Логин пользователя",
    "email": "Email пользователя",
    "password": "Хэш пароля",
    "first_name": "Имя пользователя",
    "last_name": "Фамилия пользователя",
    "is_staff": "Признак доступа к staff-части",
    "is_superuser": "Признак суперпользователя",
    "is_active": "Признак активности",
    "date_joined": "Дата регистрации",
    "last_login": "Дата и время последнего входа",
    "created_at": "Дата и время создания",
    "updated_at": "Дата и время обновления",
    "created_by_id": "Пользователь, создавший запись",
    "full_name": "Полное имя",
    "position": "Должность",
    "telegram_chat_id": "Telegram chat id",
    "candidate_telegram_chat_id": "Telegram chat id кандидата",
    "token": "Уникальный токен приглашения",
    "chat_id": "Telegram chat id",
    "linked_at": "Дата и время привязки чата",
    "action": "Тип действия",
    "object_type": "Тип объекта действия",
    "object_id": "Идентификатор объекта действия",
    "details": "Детали изменения",
    "workshop": "Цех/участок",
    "required_count": "Требуемое количество сотрудников",
    "reason": "Обоснование",
    "status": "Статус",
    "closed_at": "Дата и время закрытия",
    "candidate_name": "Имя кандидата",
    "phone": "Телефон кандидата",
    "interview_at": "Дата и время собеседования",
    "notes": "Комментарий",
    "hr_manager_id": "HR-менеджер",
    "approved_by_id": "Подтвердивший администратор",
    "employee_user_id": "Созданный пользователь-сотрудник",
    "date": "Дата",
    "start_time": "Время начала",
    "end_time": "Время окончания",
    "planned_staff": "План сотрудников",
    "assigned_staff": "Назначено сотрудников",
    "assigned_employee_id": "Ответственный сотрудник",
    "is_acknowledged": "Признак ознакомления сотрудника",
    "acknowledged_at": "Дата и время ознакомления",
    "name": "Наименование",
    "sort_order": "Порядок сортировки",
    "employee_id": "Сотрудник",
    "leave_type": "Тип отсутствия",
    "start_date": "Дата начала",
    "end_date": "Дата окончания",
    "note": "Комментарий",
    "document_name": "Наименование документа",
    "field_label": "Изменяемое поле профиля",
    "new_value": "Новое значение поля",
    "title": "Название",
    "description": "Описание",
    "priority": "Приоритет",
    "due_date": "Срок выполнения",
    "zone": "Зона/цех",
    "qualification_level": "Уровень квалификации",
    "granted_at": "Дата и время выдачи допуска",
    "granted_by_id": "Кто выдал допуск",
    "target_zone": "Целевая зона",
    "target_level": "Целевой уровень",
    "rationale": "Обоснование назначения",
    "planned_date": "Плановая дата",
    "employee_comment": "Комментарий сотрудника",
    "employee_certificate": "Файл сертификата",
    "employee_certificate_uploaded_at": "Дата загрузки сертификата",
    "hr_comment": "Комментарий HR",
    "actor_id": "Пользователь, выполнивший действие",
    "group_id": "Группа доступа",
    "permission_id": "Разрешение",
    "content_type_id": "Тип контента",
    "user_id": "Пользователь",
    "session_key": "Ключ сессии",
    "session_data": "Данные сессии",
    "expire_date": "Срок действия сессии",
    "app": "Имя приложения",
    "model": "Имя модели",
    "codename": "Код разрешения",
    "action_time": "Дата и время действия",
    "object_repr": "Строковое представление объекта",
    "action_flag": "Код типа действия",
    "change_message": "Описание изменений",
    "applied": "Дата применения миграции",
}


def requirement_text(not_null: bool, is_pk: bool, is_unique: bool, default: str | None) -> str:
    parts = ["Not null" if (not_null or is_pk) else "Null"]
    if is_unique:
        parts.append("unique")
    if default is not None:
        parts.append(f"default {default}")
    return ", ".join(parts)


def main() -> None:
    con = sqlite3.connect(DB_PATH)
    cur = con.cursor()
    cur.execute(
        "SELECT name FROM sqlite_master WHERE type='table' AND name NOT LIKE 'sqlite_%' ORDER BY name"
    )
    table_names = [row[0] for row in cur.fetchall()]

    rows: list[list[str]] = []
    for table_name in table_names:
        fk_map = {}
        for fk in cur.execute(f"PRAGMA foreign_key_list('{table_name}')").fetchall():
            _, _, ref_table, from_col, to_col, *_ = fk
            fk_map[from_col] = (ref_table, to_col)

        unique_cols = set()
        for idx in cur.execute(f"PRAGMA index_list('{table_name}')").fetchall():
            index_name = idx[1]
            is_unique = bool(idx[2])
            if not is_unique:
                continue
            cols = [c[2] for c in cur.execute(f"PRAGMA index_info('{index_name}')").fetchall()]
            if len(cols) == 1:
                unique_cols.add(cols[0])

        columns = cur.execute(f"PRAGMA table_info('{table_name}')").fetchall()
        for _, col_name, col_type, not_null, default, is_pk in columns:
            key = "PK" if is_pk else ("FK" if col_name in fk_map else "")
            data_type = col_type or "TEXT"
            required = requirement_text(bool(not_null), bool(is_pk), col_name in unique_cols, default)

            if col_name in fk_map:
                ref_table, ref_column = fk_map[col_name]
                description = f"Ссылка на {ref_table}.{ref_column}"
            else:
                description = FIELD_DESCRIPTIONS.get(col_name, f"Поле таблицы {table_name}")

            rows.append([table_name, key, col_name, data_type, required, description])

    # TSV (UTF-8 with BOM)
    with open(TSV_PATH, "w", encoding="utf-8-sig", newline="") as file:
        writer = csv.writer(file, delimiter="\t")
        writer.writerow(["Таблица", "Ключ", "Поле", "Тип данных", "Обязательность заполнения", "Описание"])
        writer.writerows(rows)

    # DOCX
    document = Document()
    document.add_heading("Таблица 8 - Словарь данных (одна таблица)", level=1)
    table = document.add_table(rows=1, cols=6)
    table.style = "Table Grid"
    headers = ["Таблица", "Ключ", "Поле", "Тип данных", "Обязательность заполнения", "Описание"]
    for i, header in enumerate(headers):
        table.rows[0].cells[i].text = header
    for row_data in rows:
        row = table.add_row().cells
        for i, value in enumerate(row_data):
            row[i].text = value

    document.save(DOCX_PATH)
    print(DOCX_PATH)
    print(TSV_PATH)


if __name__ == "__main__":
    main()
